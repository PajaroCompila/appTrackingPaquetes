from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent
IMG = BASE / "manual-capturas"
OUT = BASE / "Manual de usuario - Pedidos Bodega.docx"
LOGO = BASE.parent / "frontend/public/imagenes/logo-pajaro-azul.png"
AZUL, MEDIO, CLARO, GRIS, AVISO = "123568", "235A96", "EAF2FB", "54667A", "FFF4D6"

def shade(cell, color):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color); cell._tc.get_or_add_tcPr().append(shd)

def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(.55)
    sec.left_margin = sec.right_margin = Inches(.65)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Segoe UI", Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1E2D3D")
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [("Title",30,AZUL),("Heading 1",20,AZUL),("Heading 2",14,MEDIO)]:
        s=doc.styles[name]; s.font.name="Segoe UI"; s.font.size=Pt(size); s.font.bold=True
        s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.keep_with_next=True
    footer=sec.footer.paragraphs[0]
    footer.text="Pedidos Bodega  ·  Manual de usuario"
    footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRIS)

def heading(doc, number, title):
    p=doc.add_paragraph(style="Heading 1")
    r=p.add_run(f"{number}  "); r.font.color.rgb=RGBColor.from_string(MEDIO)
    p.add_run(title)
    return p

def callout(doc, title, body, color=CLARO):
    t=doc.add_table(rows=1, cols=1); c=t.cell(0,0); shade(c,color)
    p=c.paragraphs[0]; r=p.add_run(title+"\n"); r.bold=True; r.font.color.rgb=RGBColor.from_string(AZUL)
    p.add_run(body); doc.add_paragraph()

def step(doc, n, title, body):
    t=doc.add_table(rows=1, cols=2); t.autofit=False
    c1,c2=t.rows[0].cells; c1.width=Inches(.45); c2.width=Inches(6.45); shade(c1,MEDIO)
    p=c1.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(str(n)); r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(12)
    p=c2.paragraphs[0]; r=p.add_run(title+"\n"); r.bold=True; r.font.color.rgb=RGBColor.from_string(AZUL)
    p.add_run(body); doc.add_paragraph()

def screenshot(doc, name, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    p.add_run().add_picture(str(IMG/name), width=Inches(7.0))
    p=doc.add_paragraph(caption); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(GRIS)

def page(doc): doc.add_page_break()

def simple_table(doc, headers, rows):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text=h; shade(t.rows[0].cells[i],AZUL)
        for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=t.add_row().cells
        for i,value in enumerate(row): cells[i].text=value
    return t

doc=Document(); setup(doc)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(LOGO),width=Inches(3.4))
doc.add_paragraph()
p=doc.add_paragraph("Manual de usuario",style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("PEDIDOS BODEGA"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor.from_string(MEDIO)
p=doc.add_paragraph("Guía sencilla para buscar, preparar y revisar pedidos"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs: r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(GRIS)
doc.add_paragraph()
callout(doc,"¿Para quién es esta guía?","Para el personal de bodega, consulta y administración. No necesitás conocimientos de computación: seguí los pasos en el orden indicado.")
p=doc.add_paragraph("Versión: agosto de 2026\n\nTodos los derechos reservados. 2026© Almacén Pájaro Azul."); p.alignment=WD_ALIGN_PARAGRAPH.CENTER

page(doc); heading(doc,"1","Entrar al sistema")
doc.add_paragraph("Abrí el enlace que te dio el área de Sistemas. Vas a ver esta pantalla:")
screenshot(doc,"01-inicio-sesion.png","Pantalla de inicio de sesión.")
step(doc,1,"Escribí tu usuario","Usá el usuario que te entregó Sistemas.")
step(doc,2,"Escribí tu contraseña","Podés usar “Mostrar contraseña” para revisar lo escrito.")
step(doc,3,"Presioná “Iniciar sesión”","La aplicación abrirá la primera pantalla permitida para tu usuario.")
callout(doc,"Cuidá tu acceso","No compartás tu contraseña. Al terminar, usá el ícono de salida del menú izquierdo.",AVISO)

page(doc); heading(doc,"2","Conocer el menú")
doc.add_paragraph("El menú está en la franja azul del lado izquierdo. Pasá el cursor sobre un ícono y aparecerá su nombre.")
simple_table(doc,["Opción","Para qué sirve"],[
    ("Dashboard","Resumen de pedidos por sucursal. Solo para administradores."),
    ("Pedidos pendientes","Artículos que bodega debe revisar y preparar."),
    ("Pedidos despachados","Lo que bodega ya preparó y espera confirmación de entrega."),
    ("Historial","Pedidos que ya aparecen como entregados."),
    ("Usuarios","Crear y administrar accesos. Solo para administradores."),
    ("Cerrar sesión","Salir de la aplicación de forma segura."),
])
callout(doc,"Si una opción no aparece","No es un error. Cada persona ve solamente las pantallas permitidas para su trabajo.")

page(doc); heading(doc,"3","Dashboard: ver el movimiento del día")
doc.add_paragraph("Esta pantalla es para administradores. Muestra pedidos pendientes y entregados en cada sucursal.")
screenshot(doc,"02-dashboard.png","Dashboard con el resumen de las sucursales.")
step(doc,1,"Elegí las fechas","Indicá el período que querés revisar.")
step(doc,2,"Elegí una tienda","Podés ver todas o seleccionar una sola.")
step(doc,3,"Presioná “Aplicar filtros”","Los números se actualizarán con tu selección.")
callout(doc,"Sucursal no disponible","Podés seguir revisando las demás. Avisale a Sistemas cuál sucursal no cargó.",AVISO)

page(doc); heading(doc,"4","Buscar pedidos pendientes")
doc.add_paragraph("Cada artículo aparece en su propia fila para que puedas revisar código, cantidad y bodega sin confusión.")
screenshot(doc,"03-pedidos-pendientes.png","Pedidos pendientes con filtros y artículos.")
step(doc,1,"Buscá el pedido","Usá el número, las fechas o uno o varios almacenes.")
step(doc,2,"Presioná “Buscar”","La lista mostrará solamente lo que coincida con tus filtros.")
step(doc,3,"Revisá cada línea","Confirmá pedido, código, descripción, cantidad y bodega.")
callout(doc,"Tus filtros se conservan","Si entrás al detalle o cambiás de pantalla, la aplicación recuerda lo que elegiste. “Limpiar filtros” borra la selección.")

page(doc); heading(doc,"5","Seleccionar, imprimir y transferir")
step(doc,1,"Ver detalle","Abre el pedido completo. Al regresar, tus filtros siguen guardados.")
step(doc,2,"Consultar inventario","Hacé doble clic en el código o la descripción. Verás la existencia por bodega. Menos de 10 unidades se marca en rojo suave.")
step(doc,3,"Seleccionar","Marcá el círculo de cada artículo que vas a preparar.")
step(doc,4,"Imprimir seleccionados","Genera una sola hoja con código, descripción, cantidad y bodega.")
step(doc,5,"Transferir a despachados","Revisá el total del botón. Solo se trasladan las líneas marcadas.")
callout(doc,"Antes de transferir","No seleccionés artículos de tienda si tu bodega no los va a preparar. Revisá siempre el código de bodega.",AVISO)

page(doc); heading(doc,"6","Pedidos despachados")
doc.add_paragraph("Aquí queda lo que bodega ya preparó mientras se espera la confirmación de entrega.")
screenshot(doc,"04-pedidos-despachados.png","Pedidos preparados que esperan confirmación de entrega.")
doc.add_paragraph("Cuando la aplicación detecta la entrega, el pedido desaparece de esta pantalla y pasa al historial automáticamente.")
step(doc,1,"Revisá el pedido","Confirmá número, artículos, usuario y fecha de despacho.")
step(doc,2,"Abrí “Ver detalle”","Vas a ver todas las líneas despachadas de ese pedido.")

page(doc); heading(doc,"7","Consultar el historial")
doc.add_paragraph("El historial sirve para confirmar qué pedido fue entregado, cuándo y desde qué bodega.")
screenshot(doc,"05-historial.png","Historial de pedidos entregados.")
step(doc,1,"Buscá por número","Si lo conocés, es la forma más rápida de encontrar el pedido.")
step(doc,2,"Elegí las fechas","La búsqueda usa la fecha del pedido.")
step(doc,3,"Elegí almacenes","Podés seleccionar una o varias bodegas.")
step(doc,4,"Presioná “Buscar”","Luego usá “Ver detalle” para revisar todos los artículos.")
callout(doc,"Qué significa “Entregado”","R1 ya reporta el pedido como entregado. Los pedidos cerrados no se muestran.")

page(doc); heading(doc,"8","Administrar usuarios")
doc.add_paragraph("Esta pantalla aparece solamente para administradores.")
screenshot(doc,"06-usuarios.png","Pantalla de usuarios y accesos.")
step(doc,1,"Crear usuario","Completá nombre, usuario, contraseña y rol.")
step(doc,2,"Elegir el rol","Administrador ve todo. Operador de bodega y consulta no ven el dashboard.")
step(doc,3,"Editar o desactivar","Desactivar impide el acceso sin borrar al usuario.")
step(doc,4,"Restablecer contraseña","Usalo cuando una persona olvide su contraseña.")
callout(doc,"Cada persona con su usuario","Así queda claro quién realizó cada acción.",AVISO)

page(doc); heading(doc,"9","Si algo no funciona")
simple_table(doc,["Lo que ves","Qué podés hacer"],[
    ("Usuario o contraseña incorrectos","Revisá mayúsculas y espacios. Si continúa, pedí que restablezcan tu contraseña."),
    ("No fue posible cargar la información","Esperá unos segundos y probá otra vez. Si continúa, avisale a Sistemas."),
    ("Sucursal no disponible","Seguí con las demás y reportá cuál sucursal no cargó."),
    ("Inventario no disponible","Confirmá el código y probá otra vez. Si continúa, enviá el código a Sistemas."),
    ("No encontrás un pedido","Revisá fechas y bodegas. También puede estar en Despachados o Historial."),
])
doc.add_heading("Antes de llamar a Sistemas",level=2)
for item in ["Anotá el número de pedido o código.","Decí qué pantalla estabas usando.","Contá qué botón presionaste y qué mensaje apareció.","Tomá una captura sin mostrar tu contraseña."]:
    doc.add_paragraph(item,style="List Bullet")

page(doc); heading(doc,"10","Guía rápida para bodega")
for n,(title,body) in enumerate([
    ("Entrá con tu usuario","Nunca trabajés con el acceso de otra persona."),
    ("Abrí Pedidos pendientes","Aplicá fechas y bodegas si necesitás reducir la lista."),
    ("Revisá cada artículo","Confirmá código, cantidad y bodega."),
    ("Marcá lo que vas a preparar","La selección queda guardada mientras navegás."),
    ("Imprimí si necesitás una guía","La hoja incluye solo los artículos seleccionados."),
    ("Transferí a despachados","Revisá el total antes de presionar el botón."),
    ("Consultá el historial","Ahí aparece cuando la entrega queda confirmada."),
],1): step(doc,n,title,body)
callout(doc,"Regla sencilla","Si tenés duda con un artículo, no lo transfirás todavía. Revisalo con el encargado de bodega.",AVISO)
p=doc.add_paragraph("Contacto de Sistemas: ____________________________________"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor.from_string(AZUL)

doc.save(OUT)
print(OUT)
